"""
DOCX Parser
"""

from pathlib import Path

from docx import Document as WordDocument

from indexing.parsers.base_parser import BaseParser
from models.document_model import Document


class DOCXParser(BaseParser):

    def parse(self, file_path: str) -> Document:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(file_path)

        document = WordDocument(file_path)

        paragraphs = []

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                paragraphs.append(paragraph.text)

        text = "\n".join(paragraphs)

        return Document(
            filename=path.name,
            extension=path.suffix,
            content=text,
            source_path=str(path.resolve()),
            metadata={
                "paragraphs": len(document.paragraphs)
            }
        )